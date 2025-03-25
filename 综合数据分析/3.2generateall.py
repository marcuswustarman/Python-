import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime, timedelta
from docxtpl import DocxTemplate
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
today = datetime(2025, 3, 23).date()    


# ---------------------------
# 1. 改进版数据读取与预处理
# ---------------------------
def load_and_preprocess(file_path):
    # 读取原始数据
    raw_df = pd.read_excel(file_path, sheet_name='Sheet1', header=None)
    
    # 将单列数据转换为多列
    data = raw_df.values.reshape(-1, 5)
    df = pd.DataFrame(data, columns=['user_id', 'login_time', 'level', 'payment_amount', 'device'])
    # 多功能时间转换
    def convert_time(x):
        try:
            # 先尝试解析为Excel数值日期
            return pd.to_datetime(x, unit='D', origin='1899-12-30')
        except:
            # 再尝试解析为ISO格式字符串
            return pd.to_datetime(x, format='%Y-%m-%d %H:%M:%S', errors='coerce')
    
    df['login_time'] = df['login_time'].apply(convert_time)

    # 数据清洗
    df = df.dropna(subset=['login_time'])  # 删除无效时间记录
    df['payment_amount'] = pd.to_numeric(df['payment_amount'], errors='coerce').fillna(0)

    return df


# ---------------------------
# 2. 核心指标计算（增强版）
# ---------------------------
def calculate_kpis(df):
    today = datetime(2025, 3, 23).date()    
    yesterday = today - timedelta(days=1)
    
    # 日期筛选
    df['login_date'] = df['login_time'].dt.date
    today_df = df[df['login_date'] == today]
    yesterday_df = df[df['login_date'] == yesterday]

    # 基础指标
    dau = today_df['user_id'].nunique()
    revenue = today_df['payment_amount'].sum()
    pay_users = today_df[today_df['payment_amount'] > 0]['user_id'].nunique()
    
    

    # 次日留存率计算
    num_retained_user = len(today_df['user_id'])
    num_all_users = len(today_df['user_id']) + len(yesterday_df['user_id'])
    retention_rate = num_retained_user / num_all_users if num_all_users > 0 else 0

    #print(retention_rate)
    #print(len(today_df['user_id']))
    #print(today)
    #print(df['login_date'])
    #print(len(retained_users))
    #print(len(retained_users))

    # 设备分布
    device_dist = today_df['device'].value_counts(normalize=True).to_dict()

    return {
        '日期': today.strftime('%Y-%m-%d'),
        'DAU': dau,
        '当日收入': round(revenue, 2),
        '付费用户数': pay_users,
        '付费率(%)': round(pay_users/dau*100, 2) if dau > 0 else 0,
        'ARPPU': round(revenue/pay_users, 2) if pay_users > 0 else 0,
        '次日留存率(%)': round(retention_rate, 2),
        '设备分布': device_dist
    }

# ---------------------------
# 3. 可视化图表生成（优化版）
# ---------------------------
def generate_visualizations(df, kpi_dict):
    # 设置中文字体
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # 付费金额分布
    plt.figure(figsize=(10, 6))
    paying = df[df['payment_amount'] > 0].copy()
    bins = [0, 6, 30, 98, 198, 328, 648, np.inf]
    labels = ['6元以下', '6-30元', '30-98元', '98-198元', '198-328元', '328-648元', '648+']
    paying['amount_group'] = pd.cut(paying['payment_amount'], bins=bins, labels=labels)
    payment_dist = paying.groupby('amount_group', observed=True).size()
    payment_dist.plot(kind='bar', color='#4CAF50')
    plt.title('付费金额分布（行业标准区间）')
    plt.savefig('3.2payment_dist.png')
    plt.close()
    
    # 设备分布饼图
    plt.figure(figsize=(8, 8))
    plt.pie(
        list(kpi_dict['设备分布'].values()),
        labels=list(kpi_dict['设备分布'].keys()),
        autopct='%1.1f%%',
        colors=['#FF6F61', '#6B5B95', '#88B04B']
    )
    plt.title('设备类型分布')
    plt.savefig('3.2device_dist.png')
    plt.close()

# ---------------------------
# 4. 报告生成系统（完整版）
# ---------------------------
def generate_kpi_reports(data_dict, word_path="3.2kpi_report.docx", excel_path="3.2kpi_data.xlsx"):
    """
    生成KPI报告的Word文档和Excel表格
    
    参数：
    data_dict: 包含KPI数据的字典
    word_path: 生成的Word文档路径（默认当前目录）
    excel_path: 生成的Excel文件路径（默认当前目录）
    """
    
    def _format_value(key, value):
        """统一格式化字典值"""
        if isinstance(value, dict):
            return ", ".join([f"{k}: {v:.2%}" for k, v in value.items()])
        if isinstance(value, (np.int64, np.float64)):
            value = value.item()  # 转换numpy类型为Python原生类型
        if isinstance(value, float):
            if '%' in key:
                return f"{value*100:.2f}%"
            return f"{value:.2f}"
        return value

    # 生成Word文档（修正变量名拼写错误）
    doc = Document()
    for key, value in data_dict.items():
        formatted_value = _format_value(key, value)  # 正确变量名
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(formatted_value))  # 使用正确变量名
    doc.save(word_path)

    # 生成Excel表格
    df = pd.DataFrame({
        "指标名称": data_dict.keys(),
        "指标值": [_format_value(k, v) for k, v in data_dict.items()]
    })
    df =  df.T
    df.to_excel(excel_path, index=False)
    '''
        # 基本调用（生成默认文件）
        generate_kpi_reports(your_dict)

        # 自定义文件路径
        generate_kpi_reports(
        your_dict,
        word_path="reports/daily_report.docx",
        excel_path="data/daily_data.xlsx"
        )
    '''



# def generate_report(kpi_dict):
#     # 加载模板
#     doc = DocxTemplate("report_template.docx")
    
#     # 构建上下文
#     context = {
#         **kpi_dict,
#         'payment_dist_img': 'payment_dist.png',
#         'device_dist_img': 'device_dist.png',
#         'generation_time': datetime.now().strftime('%Y-%m-%d %H:%M')
#     }
    
#     # 渲染并保存
#     doc.render(context)
#     report_name = f"3.1运营日报_{kpi_dict['日期']}.docx"
#     doc.save(report_name)
    
#     return report_name

# ---------------------------
# 5. 邮件自动发送（安全版）
# ---------------------------
def send_email_with_report(report_path):
    msg = MIMEMultipart()
    msg['Subject'] = f"{datetime.now().strftime('%m/%d')} 游戏运营日报"
    msg['From'] = 'analytics@yourcompany.com'
    msg['To'] = 'management@yourcompany.com'
    
    # HTML正文
    html = f"""
    <h3>关键指标概览</h3>
    <ul>
      <li>DAU: {kpi_dict['DAU']}</li>
      <li>收入: ¥{kpi_dict['当日收入']}</li>
      <li>付费率: {kpi_dict['付费率(%)']}%</li>
    </ul>
    <p>详细报告请查看附件。</p>
    """
    msg.attach(MIMEText(html, 'html'))
    
    # 添加附件
    with open(report_path, "rb") as f:
        part = MIMEApplication(f.read(), Name=report_path)
    part['Content-Disposition'] = f'attachment; filename="{report_path}"'
    msg.attach(part)
    
    # 发送邮件
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login('your_email@yourcompany.com', 'your_password')
        server.send_message(msg)

# ---------------------------
# 主流程控制
# ---------------------------
def main():
    try:
        # 数据加载
        df = load_and_preprocess('originaldata.xlsx')

        # 指标计算
        kpi_dict = calculate_kpis(df)

        # 可视化
        generate_visualizations(df, kpi_dict)

        print(kpi_dict)

        generate_kpi_reports(kpi_dict)
        # 发送邮件（按需启用）
        # send_email_with_report(report_path)


        # 强制检查时间字段        
        if df['login_time'].isnull().any():
            print("警告：存在无效时间数据，已自动过滤")
            df = df.dropna(subset=['login_time'])
        
        #print(f"日报生成成功：{report_path}")

    except Exception as e:
        print(f"流程执行失败：{str(e)}")
    

if __name__ == "__main__":
    main()

